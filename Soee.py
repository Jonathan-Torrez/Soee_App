import os
import time
from datetime import datetime
import webbrowser
import speech_recognition as sr
import pyttsx3
import wikipedia
import pywhatkit
import cv2
import pandas as pd
from docx import Document
from datetime import timedelta
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import tkinter as tk
from tkinter import Text, Label, Frame, Entry, Button, ttk
from PIL import Image, ImageTk
import threading
import dateparser
import subprocess
import openai

# Inicializar el motor de texto a voz
engine = pyttsx3.init()

# Variable para guardar el contacto y número de WhatsApp
contactos = {}

# Función para que el bot hable
def talk(texto):
    engine.say(texto)
    engine.runAndWait()

def introduce():
    talk("Hola, soy Soee, ¿en qué puedo ayudarte hoy?")

def despedida():
    talk("Gracias por su atención. Muy pronto nos veremos con la tercera versión.")

def escuchar():
    reconocedor = sr.Recognizer()
    reconocedor.energy_threshold = 3000  # Umbral ajustado
    reconocedor.dynamic_energy_threshold = True  # Ajuste dinámico del umbral de energía
    with sr.Microphone() as fuente:
        try:
            audio = reconocedor.listen(fuente, timeout=5, phrase_time_limit=5)  # Tiempo de espera mayor para capturar frases largas
            comando = reconocedor.recognize_google(audio, language='es-ES')
            return comando.lower()
        except sr.UnknownValueError:
            return ""
        except sr.WaitTimeoutError:
            return ""
        except Exception as e:
            print(f"Error en escuchar: {e}")
            return ""

# Inicializar el motor de texto a voz
engine = pyttsx3.init()


# Archivos para almacenar los contactos
archivo_excel = "contactos.xlsx"
archivo_txt = "contactos.txt"

# Diccionario para guardar los contactos cargados desde Excel
contactos = {}

# Función para cargar contactos desde el archivo Excel al iniciar el programa
def cargar_contactos_excel():
    if os.path.exists(archivo_excel):
        try:
            df = pd.read_excel(archivo_excel)
            for index, row in df.iterrows():
                contactos[row['Nombre'].strip().lower()] = row['Número'].strip()
            print("Contactos cargados desde Excel correctamente.")
        except Exception as e:
            print(f"Error al cargar contactos desde Excel: {e}")
    else:
        print("No se encontró el archivo Excel. Comenzando con contactos vacíos.")


# Función para agregar un contacto desde la interfaz gráfica
def agregar_contacto():
    nombre = entry_nombre.get().strip().lower()
    pais_abreviatura = combo_paises.get().strip()
    numero = entry_numero.get().strip()

    # Verificar que los campos no estén vacíos
    if nombre and numero and pais_abreviatura:
        codigo_pais = paises_codigos.get(pais_abreviatura, "")
        numero_completo = f"{codigo_pais} {numero}".replace(" ", "")
        contactos[nombre] = numero_completo
        
        # Guardar contacto en Excel y archivo de texto
        guardar_contactos(nombre, numero_completo)
        
        # Confirmar al usuario que el contacto fue agregado
        talk(f"Contacto {nombre} agregado correctamente.")
        print(contactos)  # Verifica que el contacto esté siendo agregado al diccionario
    else:
        talk("Por favor ingrese un nombre, número y código de país válidos.")


# Función para guardar contactos en un archivo Excel
def guardar_contactos(nombre, numero):
    # Crear un DataFrame con el contacto nuevo
    df = pd.DataFrame([[nombre, numero, datetime.now().strftime('%Y-%m-%d'), datetime.now().strftime('%H:%M:%S')]],
                      columns=['Nombre', 'Número', 'Fecha de Creación', 'Hora'])
    
    try:
        # Verificar si el archivo Excel existe
        if os.path.exists(archivo_excel):
            # Si el archivo existe, agregar la nueva fila sin sobrescribir los títulos
            with pd.ExcelWriter(archivo_excel, mode='a', engine='openpyxl', if_sheet_exists='overlay') as writer:
                df.to_excel(writer, index=False, header=False)
        else:
            # Si el archivo no existe, guardar el archivo con los títulos
            df.to_excel(archivo_excel, index=False)
        
        print(f"Contacto {nombre} guardado en Excel correctamente.")
    except Exception as e:
        print(f"Error al guardar en Excel: {e}")

    # Guardar contacto en archivo de texto para mayor respaldo
    try:
        with open(archivo_txt, "a") as file:
            file.write(f"Nombre: {nombre}, Número: {numero}, Fecha de Creación: {datetime.now().strftime('%Y-%m-%d')}, Hora: {datetime.now().strftime('%H:%M:%S')}\n")
        print(f"Contacto {nombre} guardado en archivo de texto correctamente.")
    except Exception as e:
        print(f"Error al guardar en archivo de texto: {e}")

# Función para abrir el archivo Excel con los contactos
def mostrar_contactos():
    if os.path.exists(archivo_excel):
        try:
            os.startfile(archivo_excel)  # Funciona en Windows
            talk("Mostrando el archivo de contactos.")
        except Exception as e:
            talk(f"Error al abrir el archivo de contactos: {e}")
            print(f"Error al abrir el archivo Excel: {e}")
    else:
        talk("El archivo de contactos no existe.")
        print("El archivo de contactos no se encuentra.")

# Función para guardar contactos tanto en un archivo Excel como en un archivo de texto
def guardar_contactos(nombre, numero):
    df = pd.DataFrame([[nombre, numero, datetime.now().strftime('%Y-%m-%d'), datetime.now().strftime('%H:%M:%S')]],
                      columns=['Nombre', 'Número', 'Fecha de Creación', 'Hora'])
    
    try:
        if os.path.exists(archivo_excel):
            with pd.ExcelWriter(archivo_excel, mode='a', engine='openpyxl', if_sheet_exists='overlay') as writer:
                df.to_excel(writer, index=False, header=False)
        else:
            df.to_excel(archivo_excel, index=False)
        
        print(f"Contacto {nombre} guardado en Excel correctamente.")
    except Exception as e:
        print(f"Error al guardar en Excel: {e}")

    try:
        with open(archivo_txt, "a") as file:
            file.write(f"Nombre: {nombre}, Número: {numero}, Fecha de Creación: {datetime.now().strftime('%Y-%m-%d')}, Hora: {datetime.now().strftime('%H:%M:%S')}\n")
        print(f"Contacto {nombre} guardado en archivo de texto correctamente.")
    except Exception as e:
        print(f"Error al guardar en archivo de texto: {e}")

# Función para enviar un mensaje de WhatsApp
def enviar_mensaje_whatsapp(nombre, mensaje):
    try:
        nombre = nombre.lower().strip()  # Convertir el nombre a minúsculas y eliminar espacios
        contactos_bajos = {k.lower(): v for k, v in contactos.items()}  # Convertir todos los nombres a minúsculas
        
        if not nombre:
            talk("No proporcionaste un nombre de contacto.")
            return
        
        if nombre in contactos_bajos:
            numero = contactos_bajos[nombre]
            talk(f"Abriendo WhatsApp para enviar un mensaje a {nombre}...")

            # Asegurarse de que el número tenga el formato correcto
            if numero.startswith("+"):
                # Aumentar el tiempo de espera para asegurar que WhatsApp Web se mantenga abierto
                pywhatkit.sendwhatmsg_instantly(numero, mensaje, wait_time=20, tab_close=True, close_time=50)
                
                # Esperar un tiempo adicional para asegurar que el mensaje sea enviado correctamente
                time.sleep(10)  # Tiempo adicional para asegurarse de que el mensaje se envíe
                talk("Mensaje enviado con éxito.")
                print(f"Mensaje enviado a {nombre}.")
            else:
                talk("Error: El número de teléfono no tiene el formato internacional correcto.")
        else:
            talk(f"No se encontró el contacto {nombre}.")
            print(f"Error: Contacto '{nombre}' no encontrado en {contactos}")
    except Exception as e:
        talk(f"Error al enviar el mensaje: {e}")
        print(f"Error al enviar mensaje: {e}")

# Función para reproducir videos en YouTube
def reproduce(Video):
    try:
        talk(f"Reproduciendo {Video}")
        pywhatkit.playonyt(Video)
    except Exception:
        talk("No pude reproducir el video.")

# Función para buscar en Wikipedia
# Función para buscar en Wikipedia y guardar en un archivo de texto
def busca(search):
    try:
        wikipedia.set_lang("es")
        wiki = wikipedia.summary(search, sentences=1)
        talk(wiki)
        
        # Guardar la búsqueda y el resultado en un archivo de texto
        with open("busquedas.txt", "a") as file:
            file.write(f"Búsqueda: {search}\nResultado: {wiki}\nFecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        print(f"Búsqueda guardada: {search}")
    except Exception:
        talk("No pude encontrar información en Wikipedia.")
        # También guardar cuando no se encuentra información
        with open("busquedas.txt", "a") as file:
            file.write(f"Búsqueda: {search}\nResultado: No se encontró información\nFecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        print(f"Error al buscar: {search}")


# Función para buscar en Google
def buscame(something):
    talk("Buscando " + something)
    webbrowser.open_new_tab(f"https://www.google.com/search?q={something}")


def transcribir_info():
    # Pedir el título del documento
    talk("Por favor, dime el título para la transcripción.")
    titulo = escuchar()

    if not titulo:
        talk("No escuché un título. Por favor intenta de nuevo.")
        return

    # Iniciar la transcripción del contenido
    talk("Ahora, por favor empieza a dictar la información que deseas transcribir.")
    texto_transcrito = escuchar()
    
    if texto_transcrito:
        talk(f"He transcrito lo siguiente: {texto_transcrito}")
        
        # Crear el documento de Word y añadir el título y la transcripción
        doc = Document()
        doc.add_heading(titulo, level=1)  # Añadir el título como encabezado de nivel 1
        doc.add_paragraph(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}: {texto_transcrito}")
        nombre_archivo = f"{titulo}.docx"

        doc.save(nombre_archivo)

        talk(f"La transcripción se ha guardado en un archivo Word llamado {nombre_archivo}.")
    else:
        talk("No se detectó ningún texto para transcribir.")

import os
import xlsxwriter

def mostrar_archivos_word_excel():
    directorio = os.getcwd()  # Obtener el directorio actual

    # Buscar archivos con extensión .docx en el directorio actual
    archivos_word = [archivo for archivo in os.listdir(directorio) if archivo.endswith(".docx")]

    if archivos_word:
        try:
            # Crear el archivo Excel
            archivo_excel = 'archivos_word.xlsx'
            workbook = xlsxwriter.Workbook(archivo_excel)
            worksheet = workbook.add_worksheet()

            # Escribir encabezados
            worksheet.write('A1', 'Lista de Archivos Word')
            worksheet.write('B1', 'Enlace para abrir')

            # Agregar cada archivo a la lista en el archivo Excel
            fila = 1
            for archivo in archivos_word:
                ruta_completa = os.path.join(directorio, archivo)
                worksheet.write(fila, 0, archivo)  # Nombre del archivo
                worksheet.write_url(fila, 1, 'file:///' + ruta_completa.replace('\\', '/'), string='Abrir archivo')  # Enlace al archivo
                fila += 1

            # Cerrar el archivo Excel
            workbook.close()

            print(f"Archivo Excel '{archivo_excel}' creado con éxito.")
            os.startfile(archivo_excel)  # Abrir el archivo Excel generado
            talk("Mostrando el archivo Excel con la lista de archivos Word.")
        except Exception as e:
            talk(f"Error al crear o abrir el archivo Excel: {e}")
            print(f"Error al crear o abrir el archivo Excel: {e}")
    else:
        talk("No se encontraron archivos Word.")
        print("No se encontraron archivos Word en el directorio actual.")




def manejar_comandos():
    while True:
        entrada_texto.configure(state="normal")
        entrada_texto.delete("1.0", tk.END)
        entrada_texto.insert(tk.END, "Escuchando...")
        entrada_texto.configure(state="disabled")
        
        comando = escuchar()  # Escuchar el comando

        if comando:  # Solo procesar si hay un comando válido
            entrada_texto.configure(state="normal")
            entrada_texto.delete("1.0", tk.END)
            entrada_texto.insert(tk.END, f"Comando recibido: {comando}")
            entrada_texto.configure(state="disabled")

            # Ejecutar comandos
            procesar_accion(comando)
        
        time.sleep(1)  # Pequeño retraso para evitar el bucle rápido

# Inicializar el motor de texto a voz
motor = pyttsx3.init()

# Variables de autenticación para Google Calendar
SCOPES = ['https://www.googleapis.com/auth/calendar']

# Autenticación de Google Calendar
def autenticar_calendario():
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.json', 'w') as token:
            token.write(creds.to_json())
    return build('calendar', 'v3', credentials=creds)
def crear_evento_calendario(titulo, inicio, fin):
    service = autenticar_calendario()  # Autenticarse y obtener el servicio de Google Calendar
    
    # Crear el evento
    event = {
        'summary': titulo,
        'start': {
            'dateTime': inicio.isoformat(),
            'timeZone': 'America/La_Paz',  # Zona horaria de Bolivia
        },
        'end': {
            'dateTime': fin.isoformat(),
            'timeZone': 'America/La_Paz',  # Zona horaria de Bolivia
        },
    }


    # Llamar a la API para crear el evento
    try:
        evento_creado = service.events().insert(calendarId='primary', body=event).execute()
        event_id = evento_creado.get('id')
        talk(f"Evento '{titulo}' creado exitosamente. ID del evento: {event_id}")
        print(f"Evento '{titulo}' creado con ID: {event_id}")
    except Exception as e:
        talk("Ocurrió un error al crear el evento.")
        print(f"Error al crear el evento: {e}")

def agregar_evento_calendario():
    talk("¿Cuál es el título del evento?")
    titulo = escuchar()
    
    if not titulo:
        talk("No escuché un título. Por favor intenta de nuevo.")
        return
    
    while True:
        talk("¿Cuándo será el evento? Por favor menciona la fecha.")
        fecha = escuchar()
        
        if not fecha:
            talk("No escuché la fecha. Por favor repítela.")
            continue
        
        talk("¿A qué hora será el evento?")
        hora = escuchar()
        
        if not hora:
            talk("No escuché la hora. Por favor repítela.")
            continue
        
        # Combinar la fecha y la hora para crear un solo string para el análisis
        fecha_hora_completa = f"{fecha} {hora}"
        
        # Reemplazar términos de hora para una mejor interpretación
        fecha_hora_completa = fecha_hora_completa.replace("de la mañana", "am").replace("de la tarde", "pm").replace("de la noche", "pm")

        # Intentar analizar la fecha y hora de manera más flexible
        fecha_inicio = dateparser.parse(fecha_hora_completa, languages=['es'], settings={'PREFER_DAY_OF_MONTH': 'first'})

        # Si no se pudo interpretar, intentamos analizar la fecha por separado
        if not fecha_inicio:
            talk("No entendí la fecha o la hora. Por favor repite la fecha y la hora en un formato claro.")
            continue

        duracion = timedelta(hours=1)  # Duración por defecto de 1 hora
        fecha_fin = fecha_inicio + duracion
        
        crear_evento_calendario(titulo, fecha_inicio, fecha_fin)
        
        # Abrir Google Calendar y mostrar el evento creado
        webbrowser.open("https://calendar.google.com/calendar/r?tab=mc&pli=1")
        break  # Salir del bucle después de crear el evento

def decir_hora():
    # Obtiene la hora y los minutos actuales
    ahora = datetime.now()
    hora = ahora.hour
    minutos = ahora.minute

    # Determina el período del día
    if 0 <= hora < 12:
        periodo = "de la mañana"
    elif 12 <= hora < 18:
        periodo = "de la tarde"
    else:
        periodo = "de la noche"

    # Ajustar la hora para el formato de 12 horas
    if hora == 0:
        hora = 12
    elif hora > 12:
        hora -= 12

    # Formato de los minutos
    minutos_texto = f"{minutos}" if minutos > 9 else f"0{minutos}"

    # Construir la frase completa
    hora_texto = f"Son las {hora} y {minutos_texto} {periodo}"
    
    # Informa la hora actual al usuario
    talk(hora_texto)
    print(hora_texto)


def decir_fecha():
    # Obtiene la fecha actual
    ahora = datetime.now()
    dia = ahora.day
    mes = ahora.strftime('%B')
    año = ahora.year

    # Convertir el mes al español
    meses = {
        'January': 'enero', 'February': 'febrero', 'March': 'marzo',
        'April': 'abril', 'May': 'mayo', 'June': 'junio',
        'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
        'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
    }
    mes_espanol = meses.get(mes, mes)
    # Construir la frase completa
    fecha_texto = f"Hoy es {dia} de {mes_espanol} del {año}" 
    # Informa la fecha actual al usuario
    talk(fecha_texto)
    print(fecha_texto)

# Función para cerrar la ventana del asistente
def cerrar_asistente():
    talk("Cerrando el asistente. ¡Hasta pronto!")
    root.quit()  # Cierra la ventana principal de la aplicación

import os  # Asegúrate de importar la librería para apagar el computador

def despedida():
    talk("Gracias por su atención. Muy pronto nos veremos con la tercera versión.")

import subprocess

proceso_deteccion = None

def iniciar_deteccion_ojos():
    global proceso_deteccion
    # Iniciar el proceso de detección de ojos en un hilo separado
    proceso_deteccion = subprocess.Popen(['python', 'deteccion_ojos.py'])

def detener_deteccion_ojos():
    global proceso_deteccion
    if proceso_deteccion:
        proceso_deteccion.terminate()  # Terminar el proceso de detección
        proceso_deteccion = None

import sys
sys.path.append('C:/Users/Jonathan/Downloads/Soee_App/chatBot.py')  # Ruta hacia el directorio donde está chatBot.py
import chatBot  # Importar el segundo archivo chatBot.py para el uso de Soee GPT

# Variable de control para activar/desactivar Soee GPT
soee_gpt_activado = False

# Función para activar Soee GPT
def activar_soee_gpt():
    global soee_gpt_activado
    soee_gpt_activado = True
    chatBot.activar_chatgpt()  # Llamar la función de activar Soee GPT desde chatBot.py

# Función para desactivar Soee GPT
def desactivar_soee_gpt():
    global soee_gpt_activado
    soee_gpt_activado = False
    chatBot.desactivar_chatgpt()  # Llamar la función de desactivar Soee GPT desde chatBot.py

# Función para procesar comandos de voz
def procesar_accion(comando):
    global soee_gpt_activado

    if "salir" in comando:
        despedida()
        sys.exit(0)  # Cerrar la aplicación
    elif "detente" in comando:
        talk("He detenido mis actividades.")
    elif "reproduce" in comando:
        music = comando.replace('reproduce', '').strip()
        reproduce(music)
    elif "busca" in comando:
        search = comando.replace('busca', '').strip()
        busca(search)
    elif "búscame" in comando:
        something = comando.replace("búscame", '').strip()
        buscame(something)
    elif "transcribe" in comando or "transcribir" in comando:
        transcribir_info()
    elif "mensaje" in comando:
        talk("Por favor, dime el nombre del contacto.")
        nombre = escuchar() 
        if not nombre:  
            talk("No escuché un nombre de contacto. Intenta de nuevo.")
            return
        talk("¿Qué mensaje deseas enviar?")
        mensaje = escuchar()
        enviar_mensaje_whatsapp(nombre, mensaje)
    elif "evento" in comando or "calendario" in comando:
        agregar_evento_calendario()
    elif "qué hora es" in comando or "dime la hora" in comando:
        decir_hora()
    elif "qué fecha es" in comando or "dime la fecha" in comando:
        decir_fecha()
    elif "ciérrate" in comando or "cerrar" in comando:
        cerrar_asistente()
    elif "despedida" in comando:
        despedida()
        os.system("shutdown /s /t 1")  # Comando para apagar el computador en Windows
    elif "actívate" in comando:  # Activar Soee GPT
        if not soee_gpt_activado:
            activar_soee_gpt()
        else:
            talk("Soee GPT ya está activado.")
    elif "desactívate" in comando:  # Desactivar Soee GPT
        if soee_gpt_activado:
            desactivar_soee_gpt()
            talk("He desactivado la funcionalidad del asistente.")
        else:
            talk("Soee GPT ya está desactivado.")
    else:
        talk("No entendí esa solicitud. Por favor intenta de nuevo.")




# Función para centrar la ventana
def centrar_ventana(ventana, ancho, alto):
    pantalla_ancho = ventana.winfo_screenwidth()
    pantalla_alto = ventana.winfo_screenheight()
    x = (pantalla_ancho // 2) - (ancho // 2)
    y = (pantalla_alto // 2) - (alto // 2)
    ventana.geometry(f'{ancho}x{alto}+{x}+{y}')

# Configuración de la ventana principal
root = tk.Tk()
root.title("Soee")
root.geometry("900x600")
root.configure(bg="black")


# Centrar la ventana al abrir
centrar_ventana(root, 1200, 1200)

# Cargar la imagen para el icono de la ventana (pequeña)
ruta_icono = "Icono_Actual.png"  # Cambia esto por la ruta de tu imagen
icono = tk.PhotoImage(file=ruta_icono)
root.iconphoto(False, icono)

# Cargar la imagen para mostrar dentro de la ventana (más grande)
imagen = Image.open(ruta_icono)
imagen = imagen.resize((100, 100))  # Ajusta el tamaño según sea necesario
imagen_tk = ImageTk.PhotoImage(imagen)

# Crear un Label para mostrar la imagen dentro de la ventana
label_imagen = Label(root, image=imagen_tk, bg="black")
label_imagen.pack(pady=10)  # Ajusta el padding según necesites


# Marco para centrar el contenido
frame = Frame(root, bg="black")
frame.pack(expand=True)

# Cargar imagen
ruta_imagen = "IASOEE.png"  # Cambia esto a la ruta de tu imagen
imagen = Image.open(ruta_imagen)
imagen = imagen.resize((400, 400))  # Cambia el tamaño según sea necesario
imagen_tk = ImageTk.PhotoImage(imagen)
label_imagen = Label(frame, image=imagen_tk, bg="black")
label_imagen.pack(pady=10)

# Área de texto
entrada_texto = Text(frame, width=40, height=5, font=("Helvetica", 12), fg="#008080", bg="black", bd=0, wrap=tk.WORD)
entrada_texto.pack(pady=20)
entrada_texto.configure(state="normal")

# Lista de abreviaturas de países y códigos telefónicos
paises_codigos = {
    "AR": "+54", "BO": "+591", "BR": "+55",
    "CL": "+56", "CO": "+57", "EC": "+593",
    "ES": "+34", "US": "+1", "PY": "+595",
    "PE": "+51", "UY": "+598"
}

# Campos de entrada para nombre y número de contacto
Label(frame, text="Nombre del contacto:", fg="#008080", bg="black", font=("Helvetica", 10)).pack(pady=5)
entry_nombre = Entry(frame, font=("Helvetica", 12), fg="#008080", bg="black", insertbackground="#008080")
entry_nombre.pack(pady=5)

# Marco para organizar el combo_paises, bandera y entrada de número de contacto
frame_contacto = Frame(frame, bg="black")
frame_contacto.pack(pady=5)

# Etiqueta de número de contacto
Label(frame_contacto, text="Número del contacto:", fg="#008080", bg="black", font=("Helvetica", 10)).pack(side="left", padx=(0, 5))

# Crear el Combobox más pequeño con abreviaturas de países
combo_paises = ttk.Combobox(frame_contacto, values=list(paises_codigos.keys()), font=("Helvetica", 12), width=5, state="readonly")
combo_paises.pack(side="left", padx=(5, 5))

# Espacio para la bandera
label_bandera = Label(frame_contacto, bg="black")
label_bandera.pack(side="left", padx=5)

# Entrada de número de contacto
entry_numero = Entry(frame_contacto, font=("Helvetica", 12), fg="#008080", bg="black", insertbackground="#008080")
entry_numero.pack(side="left", padx=(5, 0))

# Función para mostrar la bandera seleccionada
boton_agregar = Button(frame, text="Agregar contacto", command=agregar_contacto, bg="#008080", fg="white", font=("Helvetica", 12))
boton_agregar.pack(pady=10)

# Ejemplo de uso del botón "Mostrar contactos"
boton_mostrar_contactos = Button(frame, text="Mostrar contactos", command=mostrar_contactos, bg="#008080", fg="white", font=("Helvetica", 12))
boton_mostrar_contactos.pack(pady=10)

# Botón para guardar y mostrar archivos Word con enlaces en Excel
boton_mostrar_archivos = Button(frame, text="Mostrar Archivos Word", command=mostrar_archivos_word_excel, bg="#008080", fg="white", font=("Helvetica", 12))
boton_mostrar_archivos.pack(pady=10)

# Aplicar el estilo personalizado al Combobox
style = ttk.Style()
style.theme_use('default')

# Configurar los estilos de los elementos
style.configure('TCombobox', fieldbackground='black', background='black', foreground='#008080')

# Configurar la lista desplegable del Combobox (cuando se abre)
style.map('TCombobox', fieldbackground=[('readonly', 'black')],
                         foreground=[('readonly', '#008080')],
                         background=[('readonly', 'black')])

# Configurar los estilos de la lista desplegable del Combobox
style.configure('TCombobox', selectbackground="black", selectforeground="#008080")

# Personalizar el estilo de la lista desplegable
combo_paises.option_add('*TCombobox*Listbox.background', 'black')
combo_paises.option_add('*TCombobox*Listbox.foreground', '#008080')
combo_paises.option_add('*TCombobox*Listbox.selectBackground', '#008080')
combo_paises.option_add('*TCombobox*Listbox.selectForeground', 'black')


# Iniciar el hilo para escuchar comandos
hilo_asistente = threading.Thread(target=manejar_comandos)
hilo_asistente.daemon = True  # Para que termine cuando se cierra la ventana
hilo_asistente.start()

# Bucle principal
root.mainloop()
